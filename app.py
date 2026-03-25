import streamlit as st
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import math
import os

from docx import Document
from docx.shared import Inches
from scipy.io import savemat

from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from openpyxl import Workbook
from openpyxl.styles import Font
import threading
import time

from holtrop_core import holtrop_resistance_power

# ======================================================
# AUTO-SHUTDOWN MONITOR (5 MINUTES INACTIVITY)
# ======================================================
@st.cache_resource
def start_inactivity_monitor(timeout_seconds=300):
    class InactivityMonitor:
        def __init__(self):
            self.last_active = time.time()
            self.thread = threading.Thread(target=self._check_loop, daemon=True)
            self.thread.start()

        def _check_loop(self):
            while True:
                time.sleep(10)
                if time.time() - self.last_active > timeout_seconds:
                    os._exit(0)
    return InactivityMonitor()

# Initialize or retrieve the monitor, and update the timestamp on every interaction
monitor = start_inactivity_monitor(300)
monitor.last_active = time.time()

# ======================================================
# Helper function: add borders to Word tables
# ======================================================
def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement("w:tblBorders")

    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "8")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "000000")
        tblBorders.append(border)

    tblPr.append(tblBorders)

def prevent_cell_wrap(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    noWrap = OxmlElement("w:noWrap")
    tcPr.append(noWrap)

# ======================================================
# PAGE CONFIG
# ======================================================
st.set_page_config(
    page_title="GENERAL CARGO VESSEL RESISTANCE AND CARRYING CAPACITY CALCULATION TOOL",
    page_icon="🚢",
    layout="wide"
)

st.title("GENERAL CARGO VESSEL RESISTANCE AND CARRYING CAPACITY CALCULATION TOOL")

st.markdown(
    """
    **Capabilities**
    - Holtrop–Mennen Resistance & Power Calculation.  
    - Speed-based resistance curves  
    - Generating hull data for preliminary optimization framework  
    - Export of results for MATLAB/ obtain resistance report  
    """
)

# ======================================================
# TABS
# ======================================================
tab_inputs, tab_resistance, tab_capacity, tab_optimization, tab_export = st.tabs(
    ["Inputs", "Resistance Calculations", "Carrying Capacity Calculations", "Hull Form Data", "Export"]
)

# ======================================================
# TAB 1 — INPUTS
# ======================================================
with tab_inputs:
    st.subheader("Ship Particulars (Gen. Cargo Vessel)")

    st.caption(
        "Input ranges are restricted to ensure validity of the "
        "Holtrop–Mennen resistance method for general cargo ships."
    )

    # ==================================================
    # MAIN DIMENSIONAL PARAMETERS
    # ==================================================
    st.markdown("### Main Dimensional Parameters")

    c1, c2, c3, c4 = st.columns(4)

    with c1:
        L = st.number_input(
            "Length L (m)",
            min_value=80.0,
            max_value=200.0,
            value=90.0
        )

    with c2:
        B = st.number_input(
            "Breadth B (m)",
            min_value=12.0,
            max_value=32.0,
            value=15.4
        )

    with c3:
        T = st.number_input(
            "Draft T (m)",
            min_value=5.0,
            max_value=12.0,
            value=6.55
        )

    with c4:
        CB = st.number_input(
            "Block Coefficient Cb",
            min_value=0.50,
            max_value=0.85,
            value=0.65
        )

    # ---- ranges under dimensions ----
    r1, r2, r3, r4 = st.columns(4)
    r1.caption("80 – 200 m")
    r2.caption("12 – 32 m")
    r3.caption("5 – 12 m")
    r4.caption("0.50 – 0.85")

    st.markdown("---")

    # ==================================================
    # SPEED PARAMETERS
    # ==================================================
    st.markdown("### Speed Parameters")

    s1, s2 = st.columns(2)

    with s1:
        V_min = st.number_input(
            "Minimum Speed (kn)",
            min_value=5,
            max_value=30,
            value=12
        )
  
    with s2:
        V_max = st.number_input(
            "Maximum Speed (kn)",
            min_value=5,
            max_value=30,
            value=16
        )
    st.caption(
        " Applicable Speed Range: **5-30 kn**"
    )

    # ---- design speed below (same width as min/max) ----
    sd1, sd2 = st.columns(2)

    with sd1:
        V_design = st.number_input(
        "Design Speed (kn)",
        min_value=5,
        max_value=40,
        value=14
        )


        st.caption(
        "Design speed must lie between the selected minimum and maximum speeds."
        )

    st.markdown("---")

    # ==================================================
    # SAVE INPUTS
    # ==================================================
    if st.button("Save Inputs"):
        st.session_state["L"] = L
        st.session_state["B"] = B
        st.session_state["T"] = T
        st.session_state["CB"] = CB
        st.session_state["V_min"] = int(V_min)
        st.session_state["V_max"] = int(V_max)
        st.session_state["V_design"] = int(V_design)

        st.success(
            "Inputs saved. Values are within the applicable "
            "Holtrop–Mennen range for general cargo ships."
        )

# ======================================================
# TAB 2 — RESISTANCE
# ======================================================
with tab_resistance:
    st.subheader("Resistance & Power Prediction")

    if "L" not in st.session_state:
        st.warning("Please enter and save inputs in the Inputs tab first.")
    else:
        # ----------------------------------------------
        # Resistance component selection
        # ----------------------------------------------
        st.markdown("### Select Resistance Components")

        colA, colB, colC, colD = st.columns(4)

        with colA:
            include_friction = st.checkbox(
                "Frictional Resistance", value=True, disabled=True
            )

        with colB:
            include_form = st.checkbox(
                "Form Resistance", value=True, disabled=True
            )

        with colC:
            include_wave = st.checkbox(
                "Wave-Making Resistance", value=True
            )

        with colD:
            include_bulb = st.checkbox(
                "Bulbous Bow Resistance", value=True
            )
        st.caption(
            f"Wave resistance: {'ON' if include_wave else 'OFF'} | "
            f"Bulbous bow resistance: {'ON' if include_bulb else 'OFF'}"
        )


        st.info(
            "Frictional and form resistance are always included. "
            "Wave and bulbous bow resistance may be included or excluded."
        )

        # ----------------------------------------------
        # Resistance calculation over speed range
        # ----------------------------------------------
        data = []

        for V in range(
            st.session_state["V_min"],
            st.session_state["V_max"] + 1
        ):
            
            RT, EHP, DHP, Rf, Rform, Rw, C1, C2, C5, m1, m2, Rbulb, Ct = holtrop_resistance_power(
                st.session_state["L"],
                st.session_state["B"],
                st.session_state["T"],
                st.session_state["CB"],
                V
            )
           
            V_ms = V * 0.5144
            Fn = V_ms / math.sqrt(9.81 * st.session_state["L"])
            BL = st.session_state["B"] / st.session_state["L"]
            LT = st.session_state["L"] / st.session_state["T"]
            BT = st.session_state["B"] / st.session_state["T"]


            # ----------------------------------------------
            # Apply user selection (checkbox logic)
            # ----------------------------------------------
            # Bulb resistance exists ONLY if wave resistance exists
            Rw_eff = Rw if include_wave else 0.0

            if include_wave and include_bulb:
                Rbulb_eff = Rbulb
            else:
                Rbulb_eff = 0.0


            Rtotal = Rf + Rform + Rw_eff + Rbulb_eff

           # Recalculate Ct based on effective resistance
            V_ms = V * 0.5144
            S = st.session_state["L"] * (2 * st.session_state["T"] + st.session_state["B"]) * math.sqrt(st.session_state["CB"]) * (
               0.453
               + 0.4425 * st.session_state["CB"]
               - 0.2862 * st.session_state["CB"]**2
               + 0.003467 * (st.session_state["B"] / st.session_state["T"])
            )

            if V_ms > 0:
               Ct_eff = (Rtotal * 1000) / (0.5 * 1025 * V_ms**2 * S)
            else:
               Ct_eff = 0.0

        # Display logic for wave & bulb resistance
            
            data.append([
            V,
            round(Fn, 3),
            round(Rf, 3),
            round(Rform, 3),
            round(Rw, 3),
            round(C1, 4),
            round(C2, 4),
            round(C5, 4),
            round(m1, 4),
            round(m2, 4),
            round(Rbulb_eff, 3),
            round(Rtotal, 3),
            round(Ct_eff, 6),
            round(EHP, 3),
            round(DHP, 3)
            ])


        df = pd.DataFrame(
            data,
            columns=[
                "Speed (kn)",
                "Fn",
                "Rf (kN)",
                "Rform (kN)",
                "Rw (kN)",
                "C1",
                "C2",
                "C5",
                "m1",
                "m2",
                "Rbulb (kN)",
                "Rtotal (kN)",
                "Ct",
                "EHP (kW)",
                "DHP (kW)"
            ]
        )
        
        # ----------------------------------------------
        # Hull-form ratios (constant with speed)
        # ----------------------------------------------
        L_val = st.session_state["L"]
        B_val = st.session_state["B"]
        T_val = st.session_state["T"]

        ratio_data = {
           "Ratio": ["L/B", "B/T", "L/T"],
           "Value": [
             round(L_val / B_val, 3),
             round(B_val / T_val, 3),
             round(L_val / T_val, 3),
           ]
        }

        df_ratios = pd.DataFrame(ratio_data)

        st.markdown("### Hull-Form Ratios")
        st.caption("These geometric ratios are constant for the selected hull.")
        st.dataframe(df_ratios, use_container_width=True)

        # ----------------------------------------------
        # Display resistance table
        # ----------------------------------------------
        st.markdown("### Resistance & Power Table")
        
        st.caption(
            "According to the Holtrop–Mennen method, "
            "wave-making and bulbous-bow resistance are negligible for Fn ≲ 0.25."
        )
        st.dataframe(df, use_container_width=True)

        # ----------------------------------------------
        # Design-speed resistance breakdown
        # ----------------------------------------------
        st.markdown("---")
        st.markdown(f"### Design-Speed ({st.session_state['V_design']} kn) Resistance Breakdown"
        )

        df_design = df[df["Speed (kn)"] == st.session_state["V_design"]]

        if not df_design.empty:
            row = df_design.iloc[0]

            c1, c2, c3, c4, c5, c6 = st.columns(6)

            c1.metric("Rf (kN)", f"{row['Rf (kN)']:.2f}")
            c2.metric("Rform (kN)", f"{row['Rform (kN)']:.2f}")
            Rw_val = row["Rw (kN)"] if include_wave else 0.0
            Rbulb_val = row["Rbulb (kN)"] if (include_wave and include_bulb) else 0.0

            c3.metric("Rw (kN)", f"{Rw_val:.2f}")
            c4.metric("Rbulb (kN)", f"{Rbulb_val:.2f}")

            c5.metric("Rtotal (kN)", f"{row['Rtotal (kN)']:.2f}")
            c6.metric("Ct", f"{row['Ct']:.5f}")

        # ----------------------------------------------
        # Plots
        # ----------------------------------------------
        st.markdown("---")
        st.markdown("### Resistance & Power Plots")

        col1, col2 = st.columns(2)

        # -------- Total Resistance vs Speed --------
        with col1:
            fig, ax = plt.subplots(figsize=(5, 3))

            ax.plot(
                df["Speed (kn)"],
                df["Rtotal (kN)"],
                linewidth=2,
                label="Total Resistance"
            )

            ax.axvline(
                st.session_state["V_design"],
                color="red",
                linestyle="--",
                label="Design Speed"
            )

            ax.set_xlabel("Speed (kn)")
            ax.set_ylabel("Resistance (kN)")
            ax.set_title("Total Resistance vs Speed")
            ax.grid(True)
            ax.legend(fontsize=8)

            st.pyplot(fig)

            st.caption(
             "Total Resistance variation with Speed"
            )
        # -------- Power vs Speed --------
        with col2:
            fig, ax = plt.subplots(figsize=(5, 3))

            ax.plot(df["Speed (kn)"], df["EHP (kW)"], label="EHP")
            ax.plot(df["Speed (kn)"], df["DHP (kW)"], label="DHP")

            ax.axvline(
                st.session_state["V_design"],
                color="red",
                linestyle="--",
                label="Design Speed"
            )

            ax.set_xlabel("Speed (kn)")
            ax.set_ylabel("Power (kW)")
            ax.set_title("Power vs Speed")
            ax.grid(True)
            ax.legend(fontsize=8)

            st.pyplot(fig)

            st.caption(
             "Required Power variation with Speed"
            )
        # ----------------------------------------------
        # Individual resistance component plots
        # ----------------------------------------------
        st.markdown("---")
        st.markdown("### Individual Resistance Components")

        col3, col4 = st.columns(2)

        with col3:
            fig, ax = plt.subplots(figsize=(5, 3))
            ax.plot(df["Speed (kn)"], df["Rf (kN)"], marker="o")
            ax.axvline(st.session_state["V_design"], color="red", linestyle="--")
            ax.set_xlabel("Speed (kn)")
            ax.set_ylabel("Rf (kN)")
            ax.set_title("Frictional Resistance")
            ax.grid(True)
            st.pyplot(fig)

            st.caption(
             "Frictional Resistance variation with Speed"
            )

        with col4:
            fig, ax = plt.subplots(figsize=(5, 3))
            ax.plot(df["Speed (kn)"], df["Rform (kN)"], marker="o")
            ax.axvline(st.session_state["V_design"], color="red", linestyle="--")
            ax.set_xlabel("Speed (kn)")
            ax.set_ylabel("Rform (kN)")
            ax.set_title("Form Resistance")
            ax.grid(True)
            st.pyplot(fig)

            st.caption(
             "Form Resistance variation with Speed"
            )
        col5, col6 = st.columns(2)

        with col5:
            fig, ax = plt.subplots(figsize=(5, 3))
            ax.plot(df["Speed (kn)"], df["Rw (kN)"], marker="o")
            ax.axvline(st.session_state["V_design"], color="red", linestyle="--")
            ax.set_xlabel("Speed (kn)")
            ax.set_ylabel("Rw (kN)")
            ax.set_title("Wave-Making Resistance")
            ax.grid(True)
            st.pyplot(fig)

            st.caption(
             "Wave Resistance variation with Speed"
            )

        with col6:
            fig, ax = plt.subplots(figsize=(5, 3))
            ax.plot(df["Speed (kn)"], df["Rbulb (kN)"], marker="o")
            ax.axvline(st.session_state["V_design"], color="red", linestyle="--")
            ax.set_xlabel("Speed (kn)")
            ax.set_ylabel("Rbulb (kN)")
            ax.set_title("Bulbous Bow Resistance")
            ax.grid(True)
            st.pyplot(fig)

            st.caption(
             "Bulbous Resistance variation with Speed"
            )

# ======================================================
# TAB 3 — CARRYING CAPACITY
# ======================================================
with tab_capacity:
    
    if "L" not in st.session_state:
        st.warning("Please enter and save inputs in the Inputs tab first.")
    else:

        # -----------------------------
        # Basic inputs
        # -----------------------------
        L = st.session_state["L"]
        B = st.session_state["B"]
        T = st.session_state["T"]
        Cb = st.session_state["CB"]

        rho = 1025      # seawater density (kg/m³)

        # -----------------------------
        # Hydrostatic calculations
        # -----------------------------
        block_volume = L * B * T
        displacement_volume = block_volume * Cb
        displacement_mass = displacement_volume * rho / 1000   # tonnes
        displacement_kgf = displacement_mass * 1000            # kgf equivalent

        # -----------------------------
        # Geometric ratios
        # -----------------------------
        LB = L / B
        BT = B / T
        LT = L / T
        volumetric_slenderness = L / (displacement_volume ** (1/3))

        # -----------------------------
        # Consolidated Hydrostatic Table
        # -----------------------------
        st.markdown("### Hydrostatic & Carrying Capacity Parameters")

        df_capacity = pd.DataFrame(
            {
                "Parameter": [
                    "Block Volume (L × B × T) [m³]",
                    "Block Coefficient (Cb)",
                    "Displacement Volume (∇) [m³]",
                    "Displacement Mass (Δ) [tonnes]",
                    "Displacement (kgf)",
                    "L/B Ratio",
                    "B/T Ratio",
                    "L/T Ratio",
                    "L / ∇¹ᐟ³ (Volumetric Slenderness)"
                ],
                "Value": [
                    f"{block_volume:,.2f}",
                    f"{Cb:.3f}",
                    f"{displacement_volume:,.2f}",
                    f"{displacement_mass:,.2f}",
                    f"{displacement_kgf:,.0f}",
                    f"{LB:.3f}",
                    f"{BT:.3f}",
                    f"{LT:.3f}",
                    f"{volumetric_slenderness:.3f}"
                ]
            }
        )

        st.dataframe(df_capacity, use_container_width=True)

        st.caption(
            "Hydrostatic properties derived from baseline hull dimensions. "
            "These parameters describe displacement capacity and geometric proportions "
            "relevant to carrying capacity assessment."
        )

# ======================================================
# TAB 4 — HULL-FORM DATA GENERATION
# ======================================================
with tab_optimization:

    st.subheader("Hull-Form Design Space Exploration")

    if "L" not in st.session_state:
        st.warning("Please enter and save inputs first.")
    else:

        st.markdown(
        """
        Baseline hull dimensions from Tab 1 are varied within
        user-defined percentage ranges to generate alternative
        hull-form combinations for resistance and capacity evaluation.
        """
        )

        # -------------------------------------------------
        # Baseline values
        # -------------------------------------------------
        L0 = st.session_state["L"]
        B0 = st.session_state["B"]
        T0 = st.session_state["T"]
        CB0 = st.session_state["CB"]
        Vd = st.session_state["V_design"]

        # -------------------------------------------------
        # Percentage variation sliders
        # -------------------------------------------------
        st.markdown("### Percentage Variation Around Baseline (± %)")

        c1, c2, c3, c4 = st.columns(4)

        with c1:
            pL = st.slider("L variation (%)", 0, 30, 10)

        with c2:
            pB = st.slider("B variation (%)", 0, 30, 10)

        with c3:
            pT = st.slider("T variation (%)", 0, 30, 10)

        with c4:
            pCB = st.slider("Cb variation (%)", 0, 30, 10)

        # -------------------------------------------------
        # Step sizes
        # -------------------------------------------------
        st.markdown("### Discretization Step Sizes")

        s1, s2, s3, s4 = st.columns(4)

        with s1:
            L_step = st.number_input("L step (m)", value=5.0)

        with s2:
            B_step = st.number_input("B step (m)", value=1.0)

        with s3:
            T_step = st.number_input("T step (m)", value=0.5)

        with s4:
            CB_step = st.number_input("Cb step", value=0.02)

        # -------------------------------------------------
        # Generate dataset
        # -------------------------------------------------
        if st.button("Generate Hull-Form Dataset"):

            # Baseline reference values
            RT0, EHP0, DHP0, Rf0, Rform0, Rw0, C10, C20, C50, m10, m20, Rbulb0, Ct0 = \
                holtrop_resistance_power(L0, B0, T0, CB0, Vd)

            disp_volume0 = L0 * B0 * T0 * CB0
            disp_mass0 = disp_volume0 * 1025 / 1000

            # Parameter ranges
            L_range = np.arange(L0*(1-pL/100), L0*(1+pL/100)+L_step, L_step)
            B_range = np.arange(B0*(1-pB/100), B0*(1+pB/100)+B_step, B_step)
            T_range = np.arange(T0*(1-pT/100), T0*(1+pT/100)+T_step, T_step)
            CB_range = np.arange(CB0*(1-pCB/100), CB0*(1+pCB/100)+CB_step, CB_step)

            results = []

            for L in L_range:
                for B in B_range:
                    for T in T_range:
                        for CB in CB_range:

                            RT, EHP, DHP, Rf, Rform, Rw, C1, C2, C5, m1, m2, Rbulb, Ct = \
                                holtrop_resistance_power(L, B, T, CB, Vd)

                            disp_volume = L * B * T * CB
                            disp_mass = disp_volume * 1025 / 1000

                            # Percentage change calculations
                            delta_Ct = ((Ct - Ct0) / Ct0 * 100) if Ct0 != 0 else 0.0
                            delta_disp = ((disp_mass - disp_mass0) / disp_mass0 * 100) if disp_mass0 != 0 else 0.0

                            results.append([
                                L, B, T, CB,
                                RT,
                                Ct,
                                delta_Ct,
                                disp_volume,
                                disp_mass,
                                delta_disp
                            ])

            df_gen = pd.DataFrame(
                results,
                columns=[
                    "L (m)",
                    "B (m)",
                    "T (m)",
                    "Cb",
                    "RT (kN)",
                    "Ct",
                    "δCt (%)",
                    "∇ (m³)",
                    "Δ (tonnes)",
                    "δΔ (%)"
                ]
            )

            st.session_state["df_generated"] = df_gen

            st.markdown("### Generated Hull-Form Combinations")
            st.dataframe(df_gen, use_container_width=True)

            st.success(f"{len(df_gen)} hull variants generated.")

            # -------------------------------------------------
            # Sensitivity Plots
            # -------------------------------------------------
            st.markdown("---")
            st.markdown("### Sensitivity of Ct to Individual Parameters")

            col1, col2 = st.columns(2)

            # L vs δCt
            with col1:
                df_L = df_gen[
                    np.isclose(df_gen["B (m)"], B0) &
                    np.isclose(df_gen["T (m)"], T0) &
                    np.isclose(df_gen["Cb"], CB0)
                ]


                fig, ax = plt.subplots()
                ax.plot(df_L["L (m)"], df_L["δCt (%)"], marker="o")
                ax.axhline(0, linestyle="--")
                ax.set_xlabel("Length L (m)")
                ax.set_ylabel("δCt (%)")
                ax.set_title("L vs Change in Ct")
                ax.grid(True)
                st.pyplot(fig)

            # B vs δCt
            with col2:
                df_B = df_gen[
                    np.isclose(df_gen["L (m)"], L0) &
                    np.isclose(df_gen["T (m)"], T0) &
                    np.isclose(df_gen["Cb"], CB0)
                ]


                fig, ax = plt.subplots()
                ax.plot(df_B["B (m)"], df_B["δCt (%)"], marker="o")
                ax.axhline(0, linestyle="--")
                ax.set_xlabel("Breadth B (m)")
                ax.set_ylabel("δCt (%)")
                ax.set_title("B vs Change in Ct")
                ax.grid(True)
                st.pyplot(fig)

            col3, col4 = st.columns(2)

            # T vs δCt
            with col3:
                df_T = df_gen[
                    np.isclose(df_gen["L (m)"], L0) &
                    np.isclose(df_gen["B (m)"], B0) &
                    np.isclose(df_gen["Cb"], CB0)
                ]

                fig, ax = plt.subplots()
                ax.plot(df_T["T (m)"], df_T["δCt (%)"], marker="o")
                ax.axhline(0, linestyle="--")
                ax.set_xlabel("Draft T (m)")
                ax.set_ylabel("δCt (%)")
                ax.set_title("T vs Change in Ct")
                ax.grid(True)
                st.pyplot(fig)

            # Cb vs δCt
            with col4:
                df_CB = df_gen[
                    np.isclose(df_gen["L (m)"], L0) &
                    np.isclose(df_gen["B (m)"], B0) &
                    np.isclose(df_gen["T (m)"], T0)
                ]


                fig, ax = plt.subplots()
                ax.plot(df_CB["Cb"], df_CB["δCt (%)"], marker="o")
                ax.axhline(0, linestyle="--")
                ax.set_xlabel("Block Coefficient Cb")
                ax.set_ylabel("δCt (%)")
                ax.set_title("Cb vs Change in Ct")
                ax.grid(True)
                st.pyplot(fig)


# ======================================================
# TAB 5 — EXPORT
# ======================================================
with tab_export:
    st.subheader("Export Results")

    if "L" not in st.session_state:
        st.warning("Please run resistance and optimization first.")
    else:
        st.markdown(
            """
            **Available Downloads**
            - MATLAB `.mat` file (for numerical post-processing)
            - Word `.docx` report (for analysis)
            """
        )

        # ==================================================
        # MATLAB EXPORT
        # ==================================================
 
        if st.button("Export Design-Speed Dataset (MATLAB)"):

            L0 = st.session_state["L"]
            B0 = st.session_state["B"]
            T0 = st.session_state["T"]
            CB0 = st.session_state["CB"]
            
            # Default bounds (±10%) and step sizes for MATLAB dataset generation
            L_min, L_max, L_step = L0 * 0.9, L0 * 1.1, 5.0
            B_min, B_max, B_step = B0 * 0.9, B0 * 1.1, 1.0
            T_min, T_max, T_step = T0 * 0.9, T0 * 1.1, 0.5
            CB_min, CB_max, CB_step = CB0 * 0.9, CB0 * 1.1, 0.02

            L_vals = np.arange(L_min, L_max + L_step, L_step)
            B_vals = np.arange(B_min, B_max + B_step, B_step)
            T_vals = np.arange(T_min, T_max + T_step, T_step)
            Cb_vals = np.arange(CB_min, CB_max + CB_step, CB_step)

            L_list, B_list, T_list, Cb_list, RT_list = [], [], [], [], []

            for L in L_vals:
                for B in B_vals:
                    for T in T_vals:
                        for Cb in Cb_vals:

                            # ------------------------------
                            # FEASIBILITY CONSTRAINTS
                            # ------------------------------
                            if not (5.0 <= L / B <= 9.0):
                                continue
                            if not (2.0 <= B / T <= 4.0):
                                continue
                            if not (0.55 <= Cb <= 0.85):
                                continue

                            # ------------------------------
                            # HOLTROP AT DESIGN SPEED ONLY
                            # ------------------------------
                            RT, *_ = holtrop_resistance_power(
                                L, B, T, Cb, V_design
                            )

                            L_list.append(L)
                            B_list.append(B)
                            T_list.append(T)
                            Cb_list.append(Cb)
                            RT_list.append(RT)

            # ==========================================
            # EXPORT MATLAB FILE
            # ==========================================
            savemat(
                "holtrop_designspeed_dataset.mat",
                {
                    "L": np.array(L_list),
                    "B": np.array(B_list),
                    "T": np.array(T_list),
                    "Cb": np.array(Cb_list),
                    "RT_kN": np.array(RT_list),
                }
            )

            st.success(
                f"Design-speed dataset exported "
                f"({len(RT_list)} hull variants) → holtrop_designspeed_dataset.mat"
            )


        # ==================================================
        # WORD REPORT EXPORT
        # ==================================================
        if st.button("Download Resistance Report"):

            doc = Document()
            doc.add_heading(
                "Holtrop–Mennen Resistance Analysis Report",
                level=1
            )

            # --------------------------------------------------
            # 1. Ship particulars
            # --------------------------------------------------
            doc.add_heading("1. Ship Particulars", level=2)
            p = doc.add_paragraph()
            p.add_run(f"Length (L): {st.session_state['L']} m\n")
            p.add_run(f"Breadth (B): {st.session_state['B']} m\n")
            p.add_run(f"Draft (T): {st.session_state['T']} m\n")
            p.add_run(f"Block Coefficient (Cb): {st.session_state['CB']}\n")
            p.add_run(f"Design Speed: {st.session_state['V_design']} kn\n")

            # --------------------------------------------------
            # 2. Resistance results table
            # --------------------------------------------------
            doc.add_heading("2. Resistance Results", level=2)

            headers = [
                "Speed [kn]",
                "Rf [kN]",
                "Rform [kN]",
                "Rw [kN]",
                "Rbulb [kN]",
                "RT [kN]",
                "EHP [kW]",
                "DHP [kW]",
            ]

            table = doc.add_table(rows=1, cols=len(headers))
            set_table_borders(table)

            hdr_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                hdr_cells[i].text = header
                hdr_cells[i].paragraphs[0].runs[0].bold = True
                prevent_cell_wrap(hdr_cells[i])

            for _, row in df.iterrows():
                row_cells = table.add_row().cells
                row_cells[0].text = f"{row['Speed (kn)']:.1f}"
                row_cells[1].text = f"{row['Rf (kN)']:.3f}"
                row_cells[2].text = f"{row['Rform (kN)']:.3f}"
                row_cells[3].text = f"{row['Rw (kN)']:.3f}"
                row_cells[4].text = f"{row['Rbulb (kN)']:.3f}"
                row_cells[5].text = f"{row['Rtotal (kN)']:.3f}"
                row_cells[6].text = f"{row['EHP (kW)']:.3f}"
                row_cells[7].text = f"{row['DHP (kW)']:.3f}"

            # --------------------------------------------------
            # 3. Resistance and Power Plots
            # --------------------------------------------------
            doc.add_heading("3. Resistance and Power Plots", level=2)

            plot_dir = "report_plots"
            os.makedirs(plot_dir, exist_ok=True)

            def save_plot(x, y, title, ylabel, fname):
                fig, ax = plt.subplots()
                ax.plot(x, y, linewidth=2)
                ax.set_xlabel("Speed (kn)")
                ax.set_ylabel(ylabel)
                ax.set_title(title)
                ax.grid(True)
                path = os.path.join(plot_dir, fname)
                fig.savefig(path, dpi=300, bbox_inches="tight")
                plt.close(fig)
                doc.add_paragraph(title)
                doc.add_picture(path, width=Inches(5))

            save_plot(df["Speed (kn)"], df["Rtotal (kN)"],
                      "Total Resistance vs Speed", "Resistance (kN)", "rt.png")

            save_plot(df["Speed (kn)"], df["EHP (kW)"],
                      "Effective Power vs Speed", "Power (kW)", "ehp.png")

            save_plot(df["Speed (kn)"], df["DHP (kW)"],
                      "Delivered Power vs Speed", "Power (kW)", "dhp.png")

            save_plot(df["Speed (kn)"], df["Rf (kN)"],
                      "Frictional Resistance vs Speed", "Rf (kN)", "rf.png")

            save_plot(df["Speed (kn)"], df["Rform (kN)"],
                      "Form Resistance vs Speed", "Rform (kN)", "rform.png")

            save_plot(df["Speed (kn)"], df["Rw (kN)"],
                      "Wave Resistance vs Speed", "Rw (kN)", "rw.png")

            save_plot(df["Speed (kn)"], df["Rbulb (kN)"],
                      "Bulbous Bow Resistance vs Speed", "Rbulb (kN)", "rbulb.png")

            # --------------------------------------------------
            # 4. Optimization results
            # --------------------------------------------------
            if "df_opt" in st.session_state:
                df_opt = st.session_state["df_opt"]

                doc.add_heading("4. Optimization Results", level=2)

                table2 = doc.add_table(rows=1, cols=len(df_opt.columns))
                set_table_borders(table2)

                hdr_cells = table2.rows[0].cells
                for i, col in enumerate(df_opt.columns):
                    hdr_cells[i].text = col
                    hdr_cells[i].paragraphs[0].runs[0].bold = True
                    prevent_cell_wrap(hdr_cells[i])

                for _, row in df_opt.iterrows():
                    row_cells = table2.add_row().cells
                    for i, val in enumerate(row):
                        row_cells[i].text = (
                            f"{val:.3f}" if isinstance(val, float) else str(val)
                        )

            # --------------------------------------------------
            # Save report
            # --------------------------------------------------
            doc.save("Holtrop_Resistance_Report.docx")
            st.success("Report Downloaded: Holtrop_Resistance_Report.docx")


        # ==================================================
        # EXCEL EXPORT
        # ==================================================
        if st.button("Download Excel Workbook (.xlsx)"):

            wb = Workbook()

            # ==============================
            # 1️⃣ MAINS SHEET
            # ==============================
            ws_main = wb.active
            ws_main.title = "Mains"

            ws_main["A1"] = "BASELINE MODEL PARAMETERS"
            ws_main["A1"].font = Font(bold=True)

            ws_main["A3"] = "Length (L)"
            ws_main["B3"] = st.session_state["L"]

            ws_main["A4"] = "Breadth (B)"
            ws_main["B4"] = st.session_state["B"]

            ws_main["A5"] = "Draft (T)"
            ws_main["B5"] = st.session_state["T"]

            ws_main["A6"] = "Block Coefficient (Cb)"
            ws_main["B6"] = st.session_state["CB"]

            ws_main["A7"] = "Design Speed (kn)"
            ws_main["B7"] = st.session_state["V_design"]

            # ==============================
            # 2️⃣ DIMENSIONAL SHEET
            # ==============================
            ws_dim = wb.create_sheet("Dimensional")

            if "df_generated" in st.session_state:
                df_gen = st.session_state["df_generated"]

                for col_idx, col_name in enumerate(df_gen.columns, 1):
                    ws_dim.cell(row=1, column=col_idx).value = col_name
                    ws_dim.cell(row=1, column=col_idx).font = Font(bold=True)

                for row_idx, row in enumerate(df_gen.values, 2):
                    for col_idx, value in enumerate(row, 1):
                        ws_dim.cell(row=row_idx, column=col_idx).value = float(value)

            # ==============================
            # 3️⃣ HYDROSTATICS SHEET
            # ==============================
            ws_hydro = wb.create_sheet("Hydrostatics")

            ws_hydro["A1"] = "Hydrostatic Properties"
            ws_hydro["A1"].font = Font(bold=True)

            L = st.session_state["L"]
            B = st.session_state["B"]
            T = st.session_state["T"]
            CB = st.session_state["CB"]

            disp_volume = L * B * T * CB
            disp_mass = disp_volume * 1025 / 1000

            ws_hydro["A3"] = "Displacement Volume (m³)"
            ws_hydro["B3"] = disp_volume

            ws_hydro["A4"] = "Displacement Mass (tonnes)"
            ws_hydro["B4"] = disp_mass

            ws_hydro["A6"] = "L/B"
            ws_hydro["B6"] = L / B

            ws_hydro["A7"] = "B/T"
            ws_hydro["B7"] = B / T

            ws_hydro["A8"] = "L/T"
            ws_hydro["B8"] = L / T

            # ==============================
            # 4️⃣ RESISTANCE SHEET
            # ==============================
            ws_res = wb.create_sheet("Resistance")

            if "df_resistance" in st.session_state:
                df_res = st.session_state["df_resistance"]

                for col_idx, col_name in enumerate(df_res.columns, 1):
                    ws_res.cell(row=1, column=col_idx).value = col_name
                    ws_res.cell(row=1, column=col_idx).font = Font(bold=True)

                for row_idx, row in enumerate(df_res.values, 2):
                    for col_idx, value in enumerate(row, 1):
                        ws_res.cell(row=row_idx, column=col_idx).value = float(value)

            # ==============================
            # SAVE FILE
            # ==============================
            file_name = "Holtrop_Complete_Results.xlsx"
            wb.save(file_name)

            st.success("Excel Workbook Downloaded: Holtrop_Complete_Results.xlsx")
